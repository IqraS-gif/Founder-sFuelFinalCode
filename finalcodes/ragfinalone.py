
# import streamlit as st
# import os
# from dotenv import load_dotenv
# import io
# import google.generativeai as genai
# from google.oauth2 import service_account
# from googleapiclient.discovery import build
# from googleapiclient.http import MediaIoBaseDownload
# import PyPDF2
# import docx
# import pinecone
# import numpy as np
# from sentence_transformers import SentenceTransformer
# import uuid
# import time
# from datetime import datetime
# import json
# import requests
# from serpapi import GoogleSearch

# import sqlite3
# import re
# import unicodedata
# import hashlib
# import functools
# from typing import List, Dict, Any, Optional
# import logging
# from bs4 import BeautifulSoup
# from urllib.parse import urlparse
# import pyttsx3
# import threading
# import queue

# # Configure logging
# logging.basicConfig(level=logging.WARNING)
# logger = logging.getLogger(__name__)

# # Load environment variables
# load_dotenv()

# class OptimizedStartupRAGEvaluator:
#     def __init__(self):
#         self.embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
#         self.embedding_dimension = 384
#         self.gemini_model = None
#         self.pinecone_index = None
#         self.conversation_memory = []
#         self.user_session_id = str(uuid.uuid4())
#         self.api_call_count = 0
#         self.cache = {}
#         # Fixed TTS implementation
#         self.tts_queue = queue.Queue()
#         self.tts_thread_running = False
#         self.tts_engine = None
#         self.setup_models()
#         self.setup_database()
#         self.setup_tts()
        
#     def setup_tts(self):
#         """Initialize text-to-speech engine with thread safety"""
#         try:
#             # Initialize TTS engine once
#             self.tts_engine = pyttsx3.init()
#             # Set properties for better speech quality
#             self.tts_engine.setProperty('rate', 180)  # Speed
#             self.tts_engine.setProperty('volume', 0.8)  # Volume
#             st.session_state.tts_enabled = True
            
#             # Start TTS worker thread
#             self.start_tts_worker()
            
#         except Exception as e:
#             logger.error(f"TTS setup error: {e}")
#             st.session_state.tts_enabled = False
    
#     def start_tts_worker(self):
#         """Start the TTS worker thread"""
#         if not self.tts_thread_running:
#             self.tts_thread_running = True
#             worker_thread = threading.Thread(target=self.tts_worker, daemon=True)
#             worker_thread.start()
    
#     def tts_worker(self):
#         """TTS worker thread that processes speech requests"""
#         while self.tts_thread_running:
#             try:
#                 # Get text from queue with timeout
#                 text = self.tts_queue.get(timeout=1)
#                 if text is None:  # Shutdown signal
#                     break
                
#                 # Clean text for TTS
#                 clean_text = re.sub(r'[#*\-•]', '', text)
#                 clean_text = re.sub(r'\s+', ' ', clean_text.strip())
                
#                 # Limit text length for TTS
#                 if len(clean_text) > 300:
#                     clean_text = clean_text[:300] + "..."
                
#                 # Speak the text (this is thread-safe)
#                 if self.tts_engine and clean_text.strip():
#                     self.tts_engine.say(clean_text)
#                     self.tts_engine.runAndWait()
                
#                 self.tts_queue.task_done()
                
#             except queue.Empty:
#                 continue
#             except Exception as e:
#                 logger.error(f"TTS worker error: {e}")
#                 continue
    
#     def text_to_speech(self, text: str):
#         """Queue text for speech synthesis (thread-safe)"""
#         try:
#             if not st.session_state.get('tts_enabled', False) or not self.tts_engine:
#                 return
            
#             # Add text to queue for processing
#             self.tts_queue.put(text)
                
#         except Exception as e:
#             logger.error(f"TTS queue error: {e}")
    
#     def auto_play_tts(self, text: str):
#         """Safe auto-play TTS for chat responses"""
#         self.text_to_speech(text)

#     def setup_models(self):
#         """Initialize Gemini and Pinecone"""
#         try:
#             # Configure Gemini
#             api_key = os.getenv('GEMINI_API_KEY')
#             if not api_key:
#                 st.error("GEMINI_API_KEY not found")
#                 return
#             genai.configure(api_key=api_key)
#             self.gemini_model = genai.GenerativeModel('models/gemini-2.5-flash')
            
#             # Initialize Pinecone
#             pinecone_api_key = os.getenv('PINECONE_API_KEY')
#             if not pinecone_api_key:
#                 st.error("PINECONE_API_KEY not found")
#                 return
                
#             pc = pinecone.Pinecone(api_key=pinecone_api_key)
#             index_name = "startwise-rag-knowledge"
            
#             try:
#                 pc.describe_index(index_name)
#             except:
#                 pc.create_index(
#                     name=index_name,
#                     dimension=self.embedding_dimension,
#                     metric='cosine',
#                     spec=pinecone.ServerlessSpec(cloud='aws', region='us-east-1')
#                 )
#                 time.sleep(10)
            
#             self.pinecone_index = pc.Index(index_name)
            
#         except Exception as e:
#             st.error(f"Setup error: {e}")

#     def setup_database(self):
#         """Enhanced database setup with pitch deck storage"""
#         try:
#             conn = sqlite3.connect('startup_evaluations.db', check_same_thread=False)
#             cursor = conn.cursor()
            
#             cursor.execute('''
#                 CREATE TABLE IF NOT EXISTS evaluations (
#                     id INTEGER PRIMARY KEY AUTOINCREMENT,
#                     session_id TEXT,
#                     startup_name TEXT,
#                     startup_idea TEXT,
#                     evaluation_results TEXT,
#                     pitch_deck_content TEXT,
#                     uploaded_files TEXT,
#                     timestamp DATETIME DEFAULT CURRENT_TIMESTAMP
#                 )
#             ''')
            
#             cursor.execute('''
#                 CREATE TABLE IF NOT EXISTS conversations (
#                     id INTEGER PRIMARY KEY AUTOINCREMENT,
#                     session_id TEXT,
#                     user_message TEXT,
#                     assistant_response TEXT,
#                     context_used INTEGER,
#                     timestamp DATETIME DEFAULT CURRENT_TIMESTAMP
#                 )
#             ''')
            
#             conn.commit()
#             conn.close()
#         except Exception as e:
#             st.error(f"Database setup error: {e}")
    
#     # Add this new method inside the OptimizedStartupRAGEvaluator class

#     def check_knowledge_base_status(self) -> bool:
#         """Check if the Pinecone index has vectors."""
#         try:
#             if not self.pinecone_index:
#                 st.warning("Pinecone index not initialized.")
#                 return False
            
#             # describe_index_stats() is the key function to check the state
#             stats = self.pinecone_index.describe_index_stats()
#             vector_count = stats.get('total_vector_count', 0)
            
#             logger.info(f"Pinecone index status: {vector_count} vectors found.")
            
#             # If there are vectors, we assume the KB is loaded.
#             return vector_count > 0
            
#         except Exception as e:
#             st.error(f"Failed to check knowledge base status: {e}")
#             return False

#     def clean_vector_id(self, vector_id: str) -> str:
#         """Clean vector ID for ASCII compatibility"""
#         cleaned = vector_id.replace('—', '-').replace('–', '-')
#         cleaned = unicodedata.normalize('NFKD', cleaned)
#         cleaned = cleaned.encode('ascii', 'ignore').decode('ascii')
#         cleaned = re.sub(r'[^\w\-]', '_', cleaned)
#         return cleaned.strip('_-')

#     @functools.lru_cache(maxsize=200)
#     def get_embedding(self, text: str):
#         """Get embedding with caching"""
#         return self.embedding_model.encode([text])[0]

#     def extract_text(self, file_content: bytes, file_name: str, mime_type: str) -> str:
#         """Enhanced text extraction with BeautifulSoup"""
#         try:
#             if mime_type == 'application/pdf':
#                 reader = PyPDF2.PdfReader(io.BytesIO(file_content))
#                 text = "\n".join(page.extract_text() for page in reader.pages)
                
#             elif 'wordprocessingml' in mime_type:
#                 doc = docx.Document(io.BytesIO(file_content))
#                 text = "\n".join(para.text for para in doc.paragraphs)
                
#             elif mime_type == 'text/html':
#                 soup = BeautifulSoup(file_content, 'html.parser')
#                 for script in soup(["script", "style"]):
#                     script.decompose()
#                 text = soup.get_text()
                
#             else:
#                 text = file_content.decode('utf-8', errors='ignore')
            
#             return re.sub(r'\s+', ' ', text.strip())
            
#         except Exception as e:
#             logger.error(f"Text extraction error: {e}")
#             return ""

#     def process_uploaded_files(self, uploaded_files) -> str:
#         """Process uploaded pitch deck or startup documents"""
#         combined_content = ""
        
#         for uploaded_file in uploaded_files:
#             try:
#                 file_content = uploaded_file.read()
                
#                 # Determine MIME type from file extension
#                 file_ext = uploaded_file.name.lower().split('.')[-1]
#                 mime_mapping = {
#                     'pdf': 'application/pdf',
#                     'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
#                     'txt': 'text/plain',
#                     'html': 'text/html'
#                 }
                
#                 mime_type = mime_mapping.get(file_ext, 'text/plain')
                
#                 # Extract text
#                 extracted_text = self.extract_text(file_content, uploaded_file.name, mime_type)
                
#                 if extracted_text:
#                     combined_content += f"\n\n=== FROM {uploaded_file.name.upper()} ===\n"
#                     combined_content += extracted_text[:2000]  # Limit per file
                    
#             except Exception as e:
#                 st.warning(f"Could not process {uploaded_file.name}: {e}")
        
#         return combined_content

#     def web_search_with_scraping(self, query: str, max_results: int = 3) -> List[Dict]:
#         """Web search with content scraping for market intelligence"""
#         try:
#             serpapi_key = os.getenv('SERPAPI_API_KEY')
#             if not serpapi_key:
#                 return []

#             search = GoogleSearch({
#                 "q": query,
#                 "api_key": serpapi_key,
#                 "engine": "google",
#                 "num": max_results
#             })
            
#             results = search.get_dict()
#             enhanced_results = []
            
#             for result in results.get("organic_results", []):
#                 enhanced_result = {
#                     'title': result.get('title', ''),
#                     'snippet': result.get('snippet', ''),
#                     'link': result.get('link', ''),
#                     'source': 'web_search',
#                     'timestamp': datetime.now().isoformat()
#                 }
                
#                 # Scrape content if it's a relevant business site
#                 if self.should_scrape_url(result.get('link', '')):
#                     scraped_content = self.scrape_url_content(result.get('link', ''))
#                     if scraped_content:
#                         enhanced_result['full_content'] = scraped_content
                
#                 enhanced_results.append(enhanced_result)
            
#             return enhanced_results
            
#         except Exception as e:
#             logger.error(f"Web search error: {e}")
#             return []

#     def should_scrape_url(self, url: str) -> bool:
#         """Determine if URL should be scraped"""
#         if not url:
#             return False
            
#         trusted_domains = [
#             'techcrunch.com', 'crunchbase.com', 'pitchbook.com', 
#             'cbinsights.com', 'ycombinator.com', 'medium.com', 
#             'forbes.com', 'entrepreneur.com', 'venturebeat.com',
#             'bloomberg.com', 'reuters.com', 'wsj.com'
#         ]
        
#         parsed_url = urlparse(url)
#         domain = parsed_url.netloc.lower()
        
#         return any(trusted_domain in domain for trusted_domain in trusted_domains)

#     def scrape_url_content(self, url: str) -> Optional[str]:
#         """Scrape URL content using BeautifulSoup"""
#         try:
#             headers = {
#                 'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
#             }
            
#             response = requests.get(url, headers=headers, timeout=10)
#             response.raise_for_status()
            
#             soup = BeautifulSoup(response.content, 'html.parser')
            
#             # Remove unwanted elements
#             for element in soup(['script', 'style', 'nav', 'footer', 'header']):
#                 element.decompose()
            
#             # Extract main content
#             main_content = soup.find('main') or soup.find('article') or soup.find('div', class_='content')
            
#             if main_content:
#                 text = main_content.get_text(separator=' ', strip=True)
#             else:
#                 text = soup.get_text(separator=' ', strip=True)
            
#             # Clean the text
#             text = re.sub(r'\s+', ' ', text)
#             return text[:2000]  # Limit to prevent overload
            
#         except Exception as e:
#             logger.error(f"URL scraping error for {url}: {e}")
#             return None

#     def chunk_text(self, text: str, source_file: str, domain: str, chunk_size: int = 400) -> List[Dict]:
#         """Simple text chunking"""
#         words = text.split()
#         chunks = []
        
#         for i in range(0, len(words), chunk_size):
#             chunk_words = words[i:i + chunk_size]
#             chunk_text = ' '.join(chunk_words)
            
#             if len(chunk_text.strip()) > 50:
#                 chunks.append({
#                     'text': chunk_text,
#                     'source_file': source_file,
#                     'domain': domain,
#                     'chunk_index': len(chunks),
#                     'timestamp': datetime.now().isoformat()
#                 })
        
#         return chunks

#     def retrieve_context(self, query: str, top_k: int = 5) -> List[Dict]:
#         """Simple context retrieval from knowledge base"""
#         try:
#             query_embedding = self.get_embedding(query)
            
#             results = self.pinecone_index.query(
#                 vector=query_embedding.tolist(),
#                 top_k=top_k,
#                 include_metadata=True
#             )
            
#             contexts = []
#             for match in results.get('matches', []):
#                 if match['score'] > 0.25:
#                     contexts.append({
#                         'text': match['metadata']['text'],
#                         'source': match['metadata']['source_file'],
#                         'domain': match['metadata']['domain'],
#                         'similarity': match['score']
#                     })
            
#             return contexts
            
#         except Exception as e:
#             logger.error(f"Context retrieval error: {e}")
#             return []

#     def comprehensive_startup_evaluation(self, startup_idea: str, startup_name: str = "", uploaded_content: str = "") -> Dict[str, str]:
#         """SINGLE Gemini API call covering ALL 7 domains with web-enhanced context"""
        
#         st.write("🔍 Gathering knowledge base context...")
#         rag_contexts = self.retrieve_context(startup_idea, top_k=8)
        
#         st.write("🌐 Searching for current market data...")
#         # Enhanced web search for comprehensive market analysis
#         web_searches = [
#             f"{startup_idea} startup competitors market analysis",
#             f"{startup_idea} technology stack architecture",
#             f"{startup_idea} funding investment trends 2024"
#         ]
        
#         all_web_contexts = []
#         for search_query in web_searches:
#             web_results = self.web_search_with_scraping(search_query, max_results=2)
#             all_web_contexts.extend(web_results)
        
#         # Build comprehensive context with all data sources
#         context_text = ""
        
#         # Add uploaded content context (pitch deck, etc.)
#         if uploaded_content.strip():
#             context_text += "\n=== UPLOADED STARTUP DOCUMENTS ===\n"
#             context_text += uploaded_content[:2000] + "\n\n"
        
#         # Add knowledge base context
#         if rag_contexts:
#             context_text += "\n=== EXPERT KNOWLEDGE BASE ===\n"
#             for ctx in rag_contexts[:5]:
#                 context_text += f"[{ctx['domain']} - {ctx['source']}] (Relevance: {ctx['similarity']:.2f})\n"
#                 context_text += f"{ctx['text']}\n\n"
        
#         # Add web context
#         if all_web_contexts:
#             context_text += "\n=== CURRENT MARKET INTELLIGENCE ===\n"
#             for ctx in all_web_contexts:
#                 context_text += f"• {ctx['title']}\n{ctx['snippet']}\n"
#                 if ctx.get('full_content'):
#                     context_text += f"Full Content: {ctx['full_content'][:400]}...\n\n"
        
#         # SINGLE comprehensive prompt covering ALL 7 domains explicitly
#         comprehensive_prompt = f"""You are FoundersFuelAI, an expert startup advisor. Analyze this startup covering ALL 7 critical evaluation domains.

# STARTUP INFORMATION:
# - Name: {startup_name or 'Not specified'}
# - Idea: {startup_idea}

# {context_text}

# Provide a comprehensive evaluation covering these EXACT 7 domains:

# ## 1. STARTUP IDEA EVALUATION - Comprehensive Business Concept Analysis
# - Problem-solution fit analysis
# - Target market identification and size
# - Business model viability assessment
# - Value proposition strength
# - Market timing and opportunity window
# - Core concept validation score (1-10) with reasoning

# ## 2. UNIQUENESS CHECK - Market Differentiation & Competitive Analysis
# - Analyze market differentiation potential
# - Identify direct and indirect competitors
# - Assess competitive advantages and moats
# - Market positioning opportunities
# - Uniqueness score (1-10) with reasoning

# ## 3. TECH STACK RECOMMENDATION - Technology Choices & Architecture
# - Recommended technology stack (frontend, backend, database)
# - Architecture patterns and scalability considerations
# - Third-party integrations and APIs needed
# - Development complexity assessment
# - Technology risk factors

# ## 4. SIMILAR STARTUPS - Market Landscape & Competitors
# - List 5-8 direct/indirect competitors with details
# - Competitive landscape mapping
# - Market share analysis
# - Success/failure patterns in this space
# - Competitive intelligence insights

# ## 5. PITCH GENERATION - Professional Elevator Pitches for Investors
# ### 60-Second Hooking Pitch:
# [Create a compelling 60-second pitch designed to hook investor attention - focus on the problem, solution, and initial traction]

# ### In-Depth Investment Pitch:
# [Comprehensive pitch with key points to elaborate:]
# - Problem Statement & Market Pain
# - Solution & Technology Advantage
# - Market Size & Opportunity
# - Business Model & Revenue Streams
# - Competitive Landscape & Differentiation
# - Traction & Metrics
# - Team & Execution Capability
# - Financial Projections & Unit Economics
# - Funding Requirements & Use of Funds
# - Exit Strategy & ROI Potential

# ### Key Messaging Framework:
# - Value proposition statements
# - Investment thesis points
# - Risk mitigation talking points

# ## 6. IMPROVEMENT SUGGESTIONS - Enhancement Recommendations
# - Product development priorities
# - Market strategy optimizations
# - Business model improvements
# - Technical architecture enhancements
# - Go-to-market refinements

# ## 7. SUCCESS PROBABILITY & INVESTMENT METRICS
# - Overall success probability (%) with detailed reasoning
# - Investment attractiveness rating (1-10)
# - Market timing assessment
# - Scalability potential
# - Risk-reward analysis
# - Funding readiness score

# ADDITIONAL ANALYSIS:
# - Revenue model recommendations
# - Key performance indicators (KPIs)
# - Milestone roadmap (6-month, 1-year, 3-year)
# - Partnership opportunities
# - Market entry strategy

# Base your analysis on the provided expert knowledge, current market data, and uploaded documents. Be specific with numbers, examples, and actionable recommendations. Reference sources when making claims.
# **IMPORTANT: Your entire response must be strictly under 3235 characters.**"""

#         # SINGLE API CALL for entire evaluation
#         try:
#             st.write("🧠 Generating comprehensive evaluation...")
#             response = self.gemini_model.generate_content(comprehensive_prompt)
#             self.api_call_count += 1
            
#             # Parse response into structured sections
#             sections = self.parse_evaluation_response(response.text)
#             return sections
            
#         except Exception as e:
#             return {"error": f"Error generating evaluation: {e}"}

#     def parse_evaluation_response(self, response_text: str) -> Dict[str, str]:
#         """Parse the evaluation response into structured sections"""
#         sections = {
#             "idea_evaluation": "",
#             "uniqueness": "",
#             "tech_stack": "",
#             "similar_startups": "",
#             "pitch": "",
#             "improvements": "",
#             "success_metrics": "",
#             "additional": ""
#         }
        
#         try:
#             # Simple section parsing based on headers
#             current_section = "general"
#             lines = response_text.split('\n')
            
#             for line in lines:
#                 lower_line = line.lower()
                
#                 if "startup idea evaluation" in lower_line or "business concept analysis" in lower_line:
#                     current_section = "idea_evaluation"
#                 elif "uniqueness check" in lower_line or "market differentiation" in lower_line:
#                     current_section = "uniqueness"
#                 elif "tech stack" in lower_line or "technology choices" in lower_line:
#                     current_section = "tech_stack"
#                 elif "similar startups" in lower_line or "market landscape" in lower_line:
#                     current_section = "similar_startups"
#                 elif "pitch generation" in lower_line or "elevator pitch" in lower_line:
#                     current_section = "pitch"
#                 elif "improvement suggestions" in lower_line or "enhancement recommendations" in lower_line:
#                     current_section = "improvements"
#                 elif "success probability" in lower_line or "investment metrics" in lower_line:
#                     current_section = "success_metrics"
#                 elif "additional analysis" in lower_line:
#                     current_section = "additional"
#                 else:
#                     if current_section in sections:
#                         sections[current_section] += line + "\n"
#                     else:
#                         sections["additional"] += line + "\n"
            
#             # If parsing fails, put everything in additional
#             if not any(sections.values()):
#                 sections["additional"] = response_text
                
#         except Exception as e:
#             logger.error(f"Parsing error: {e}")
#             sections["additional"] = response_text
        
#         return sections

#     def enhanced_chat_with_web_context(self, user_message: str) -> tuple[str, List[Dict]]:
#         """Chat with web context but SINGLE Gemini call"""
        
#         # Get RAG context
#         rag_contexts = self.retrieve_context(user_message, top_k=5)
        
#         # Get web context for market/trend questions
#         web_contexts = []
#         market_keywords = ['market', 'competitor', 'trend', 'industry', 'funding', 'investment']
#         if any(keyword in user_message.lower() for keyword in market_keywords):
#             web_contexts = self.web_search_with_scraping(user_message, max_results=3)
        
#         # Build enhanced context
#         context_text = ""
        
#         if rag_contexts:
#             context_text += "Knowledge Base:\n"
#             for ctx in rag_contexts[:3]:
#                 context_text += f"[{ctx['domain']}] {ctx['text'][:300]}...\n\n"
        
#         if web_contexts:
#             context_text += "Current Market Data:\n"
#             for ctx in web_contexts:
#                 context_text += f"• {ctx['title']}: {ctx['snippet']}\n"
#                 if ctx.get('full_content'):
#                     context_text += f"Details: {ctx['full_content'][:200]}...\n\n"
        
#         # Add conversation history for context
#         if self.conversation_memory:
#             context_text += "\nRecent conversation:\n"
#             for turn in self.conversation_memory[-2:]:  # Last 2 turns only
#                 context_text += f"Q: {turn['user']}\nA: {turn['assistant'][:200]}...\n\n"
        
#         # SINGLE API call with all context
#         prompt = f"""You are FoundersFuelAI, a startup advisor. Answer this question using the provided context.

# {context_text}

# User Question: {user_message}

# Provide a helpful, specific answer based on the context above. Reference sources when relevant. **Your entire response must be under 3235 characters.**"""

#         try:
#             response = self.gemini_model.generate_content(prompt)
#             self.api_call_count += 1
            
#             # Save to memory
#             self.conversation_memory.append({
#                 'user': user_message,
#                 'assistant': response.text,
#                 'contexts': rag_contexts + web_contexts
#             })
            
#             # Keep memory manageable
#             if len(self.conversation_memory) > 10:
#                 self.conversation_memory = self.conversation_memory[-5:]
            
#             return response.text, rag_contexts + web_contexts
            
#         except Exception as e:
#             return f"Error: {e}", []

#     def build_knowledge_base(self, service, parent_folder_id: str) -> bool:
#         """Build knowledge base with progress tracking"""
#         try:
#             # Get domain folders
#             query = f"'{parent_folder_id}' in parents and mimeType='application/vnd.google-apps.folder'"
#             response = service.files().list(q=query, fields='files(id, name)').execute()
#             domain_folders = response.get('files', [])
            
#             if not domain_folders:
#                 st.error("No domain folders found")
#                 return False
            
#             progress_bar = st.progress(0)
#             total_vectors = 0
            
#             for folder_idx, folder in enumerate(domain_folders):
#                 domain_name = folder['name']
#                 folder_id = folder['id']
#                 namespace = f"domain_{self.clean_vector_id(domain_name)}"
                
#                 st.write(f"Processing: {domain_name}")
                
#                 files = self.get_drive_files(service, folder_id)
#                 all_chunks = []
                
#                 for file in files:
#                     file_content = self.download_file(service, file)
#                     if file_content:
#                         content = self.extract_text(file_content, file['name'], file['mimeType'])
#                         if content.strip():
#                             chunks = self.chunk_text(content, file['name'], domain_name)
#                             all_chunks.extend(chunks)
                
#                 # Upload to Pinecone
#                 if all_chunks:
#                     vectors_to_upsert = []
#                     for chunk in all_chunks:
#                         embedding = self.get_embedding(chunk['text'])
#                         vector_id = f"{self.clean_vector_id(domain_name)}_{self.clean_vector_id(chunk['source_file'])}_{chunk['chunk_index']}"
                        
#                         vectors_to_upsert.append({
#                             'id': vector_id,
#                             'values': embedding.tolist(),
#                             'metadata': {
#                                 'text': chunk['text'],
#                                 'source_file': chunk['source_file'],
#                                 'domain': chunk['domain'],
#                                 'chunk_index': chunk['chunk_index'],
#                                 'timestamp': chunk['timestamp']
#                             }
#                         })
                    
#                     # Upload in batches
#                     batch_size = 100
#                     for i in range(0, len(vectors_to_upsert), batch_size):
#                         batch = vectors_to_upsert[i:i + batch_size]
#                         self.pinecone_index.upsert(vectors=batch, namespace=namespace)
                    
#                     total_vectors += len(vectors_to_upsert)
                
#                 progress_bar.progress((folder_idx + 1) / len(domain_folders))
            
#             st.success(f"Knowledge base built: {total_vectors} vectors")
#             return True
            
#         except Exception as e:
#             st.error(f"Knowledge base build error: {e}")
#             return False

#     def get_drive_files(self, service, folder_id: str) -> List[Dict]:
#         """Get files from Google Drive folder"""
#         try:
#             query = f"'{folder_id}' in parents"
#             results = service.files().list(
#                 q=query,
#                 fields="files(id, name, mimeType)"
#             ).execute()
            
#             files = results.get('files', [])
#             return [f for f in files if f['mimeType'] in [
#                 'application/pdf',
#                 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
#                 'application/vnd.google-apps.document',
#                 'text/plain',
#                 'text/html'
#             ]]
            
#         except Exception as e:
#             logger.error(f"File retrieval error: {e}")
#             return []

#     def download_file(self, service, file: Dict) -> Optional[bytes]:
#         """Download file from Google Drive"""
#         try:
#             if file['mimeType'] == 'application/vnd.google-apps.document':
#                 request = service.files().export_media(fileId=file['id'], mimeType='text/plain')
#             else:
#                 request = service.files().get_media(fileId=file['id'])
            
#             file_content = io.BytesIO()
#             downloader = MediaIoBaseDownload(file_content, request)
            
#             done = False
#             while not done:
#                 status, done = downloader.next_chunk()
            
#             file_content.seek(0)
#             return file_content.read()
            
#         except Exception as e:
#             logger.error(f"Download error: {e}")
#             return None

#     def save_evaluation(self, startup_name: str, startup_idea: str, results: Dict[str, str], uploaded_files: str = ""):
#         """Save evaluation to database with uploaded files info"""
#         try:
#             conn = sqlite3.connect('startup_evaluations.db', check_same_thread=False)
#             cursor = conn.cursor()
#             cursor.execute('''
#                 INSERT INTO evaluations (session_id, startup_name, startup_idea, evaluation_results, uploaded_files)
#                 VALUES (?, ?, ?, ?, ?)
#             ''', (self.user_session_id, startup_name, startup_idea, json.dumps(results), uploaded_files))
#             conn.commit()
#             conn.close()
#         except Exception as e:
#             logger.error(f"Save error: {e}")

#     def save_conversation(self, user_msg: str, assistant_response: str, context_count: int):
#         """Save conversation to database"""
#         try:
#             conn = sqlite3.connect('startup_evaluations.db', check_same_thread=False)
#             cursor = conn.cursor()
#             cursor.execute('''
#                 INSERT INTO conversations (session_id, user_message, assistant_response, context_used)
#                 VALUES (?, ?, ?, ?)
#             ''', (self.user_session_id, user_msg, assistant_response, context_count))
#             conn.commit()
#             conn.close()
#         except Exception as e:
#             logger.error(f"Conversation save error: {e}")

#     def authenticate_google_drive(self):
#         """Authenticate with Google Drive"""
#         try:
#             credentials = service_account.Credentials.from_service_account_file(
#                 'service-account.json',
#                 scopes=['https://www.googleapis.com/auth/drive.readonly']
#             )
#             return build('drive', 'v3', credentials=credentials)
#         except Exception as e:
#             st.error(f"Authentication failed: {e}")
#             return None

#     def load_evaluation_history(self) -> List[tuple]:
#         """Load evaluation history"""
#         try:
#             conn = sqlite3.connect('startup_evaluations.db', check_same_thread=False)
#             cursor = conn.cursor()
#             cursor.execute('''
#                 SELECT startup_name, startup_idea, timestamp 
#                 FROM evaluations 
#                 ORDER BY timestamp DESC 
#                 LIMIT 10
#             ''')
#             results = cursor.fetchall()
#             conn.close()
#             return results
#         except:
#             return []

#     def cleanup_tts(self):
#         """Cleanup TTS resources on shutdown"""
#         try:
#             self.tts_thread_running = False
#             if hasattr(self, 'tts_queue'):
#                 self.tts_queue.put(None)  # Shutdown signal
#             if self.tts_engine:
#                 self.tts_engine.stop()
#         except Exception as e:
#             logger.error(f"TTS cleanup error: {e}")


# def main():
#     st.set_page_config(
#         page_title="Founder's Fuel - Complete Startup Evaluator", 
#         page_icon="🚀",
#         layout="wide"
#     )
    
#     # Clean, simple styling
#     st.markdown("""
#     <style>
#     .main-header {
#         text-align: center;
#         padding: 2rem;
#         background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
#         color: white;
#         border-radius: 15px;
#         margin-bottom: 2rem;
#         box-shadow: 0 8px 32px rgba(102, 126, 234, 0.3);
#     }
#     .status-success {
#         background: #d4edda;
#         color: #155724;
#         border: 2px solid #28a745;
#         border-radius: 8px;
#         text-align: center;
#         padding: 15px;
#         font-weight: bold;
#     }
#     .status-warning {
#         background: #fff3cd;
#         color: #856404;
#         border: 1px solid #ffeeba;
#         border-radius: 8px;
#         text-align: center;
#         padding: 15px;
#     }
#     .metric-card {
#         background: white;
#         padding: 1rem;
#         border-radius: 10px;
#         border: 1px solid #e2e8f0;
#         box-shadow: 0 2px 4px rgba(0, 0, 0, 0.1);
#         text-align: center;
#     }
#         [data-testid="stExpander"] {
#     border-radius: 10px;
#     border: 1px solid #e2e8f0;
#     transition: box-shadow 0.3s ease-in-out;
# }

# [data-testid="stExpander"]:hover {
#     box-shadow: 0 4px 8px rgba(0, 0, 0, 0.1);
# }
                        
# [data-testid="stExpander"] summary p {
#     font-size: 22px !important; /* Increase title font size */
#     font-weight: bold !important;
#             color:#9EFCFF !important;
# }
#     .st-expander-content div {
#         font-size: 19px !important; /* Increases the font size of the content within the expander */
#     }
                
                

#     </style>
#     """, unsafe_allow_html=True)
    
#     # Header
#     st.markdown("""
#     <div class="main-header">
#         <h1>🚀 Founder's Fuel</h1>
#         <h3>Complete Startup Evaluation Platform</h3>
#         <p>7-Domain Analysis • Web-Enhanced RAG • Single API Call Optimization</p>
#     </div>
#     """, unsafe_allow_html=True)
    
#     # Initialize evaluator
#     # --- MODIFIED CODE ---
#     # Initialize evaluator and check KB status on first run
#     if 'evaluator' not in st.session_state:
#         with st.spinner("Initializing evaluator and checking knowledge base..."):
#             st.session_state.evaluator = OptimizedStartupRAGEvaluator()
#             # Automatically check if the KB is already built in Pinecone
#             st.session_state.kb_loaded = st.session_state.evaluator.check_knowledge_base_status()

#     # Simple sidebar
#     with st.sidebar:
#         st.markdown("### 🛠️ Setup & Configuration")
        
#         # TTS Toggle for chat only
#         tts_enabled = st.toggle("🔊 Enable Chat TTS", value=st.session_state.get('tts_enabled', True))
#         st.session_state.tts_enabled = tts_enabled
        
#         if tts_enabled:
#             st.success("🔊 Chat TTS Active")
#         else:
#             st.warning("🔇 Chat TTS Disabled")
        
#         st.divider()
        
#         parent_folder_id = st.text_input(
#             "Google Drive Knowledge Folder ID", 
#             value="1fZ2p0uvCKsGtqrofcVc-SeGxq0YCmUNT"
#         )
        
#         # --- MODIFIED CODE ---
#         if st.button("🔧 Build / Rebuild Knowledge Base", type="primary"):
#             service = st.session_state.evaluator.authenticate_google_drive()
#             if service:
#                 with st.spinner("Building knowledge base..."):
#                     if st.session_state.evaluator.build_knowledge_base(service, parent_folder_id):
#                         st.session_state.kb_loaded = True
#                         st.success("Knowledge base ready!")
#                         st.rerun()
        
#         # Status
#         if st.session_state.kb_loaded:
#             st.success("✅ Knowledge Base Ready")
#         else:
#             st.warning("⚠️ Knowledge Base Not Loaded")
        
#         st.divider()
        
#         # Metrics
#         st.markdown("### 📈 Usage Metrics")
#         col1, col2 = st.columns(2)
#         with col1:
#             st.metric("Gemini Calls", st.session_state.evaluator.api_call_count)
#         with col2:
#             st.metric("Session", st.session_state.evaluator.user_session_id[-6:])

#     # Main content
#     if not st.session_state.kb_loaded:
#         st.info("Please build the knowledge base first using the sidebar to unlock all features.")
#         return

#     # Tabs - Reordered to put evaluation first
#     tab1, tab2 = st.tabs(["📊 7-Domain Evaluation", "💬 AI Chat"])
    
#     with tab1:
#         st.markdown("### 📊 Complete 7-Domain Startup Evaluation")
        
#         # Document Upload Section (moved here)
#         st.markdown("### 📁 Upload Supporting Documents")
#         uploaded_files = st.file_uploader(
#             "Upload pitch deck, business plan, or other startup documents",
#             type=['pdf', 'docx', 'txt', 'html'],
#             accept_multiple_files=True,
#             help="Upload documents to provide better context for your evaluation"
#         )
        
#         uploaded_content = ""
#         if uploaded_files:
#             st.success(f"📄 {len(uploaded_files)} file(s) uploaded!")
            
#             for file in uploaded_files:
#                 st.write(f"• **{file.name}** ({file.size:,} bytes)")
            
#             with st.spinner("Processing files..."):
#                 uploaded_content = st.session_state.evaluator.process_uploaded_files(uploaded_files)
#                 if uploaded_content:
#                     st.success("✅ Documents processed and ready for evaluation!")
        
#         st.divider()
        
#         with st.form("evaluation_form"):
#             startup_name = st.text_input("🏢 Startup Name (optional)")
#             startup_idea = st.text_area(
#                 "💡 Startup Idea & Description", 
#                 height=120,
#                 placeholder="Describe your startup concept, problem it solves, target market..."
#             )
            
#             submitted = st.form_submit_button("🚀 Generate Complete Evaluation", type="primary")
        
#         if submitted and startup_idea.strip():
#             start_time = time.time()
            
#             with st.spinner("Generating comprehensive analysis..."):
#                 evaluation_sections = st.session_state.evaluator.comprehensive_startup_evaluation(
#                     startup_idea, startup_name, uploaded_content
#                 )
            
#             if "error" not in evaluation_sections:
#                 # Save results
#                 uploaded_files_info = ", ".join([f.name for f in uploaded_files]) if uploaded_files else ""
#                 st.session_state.evaluator.save_evaluation(
#                     startup_name, startup_idea, evaluation_sections, uploaded_files_info
#                 )
                
#                 st.success("📋 Evaluation Complete!")
                
#                 # Display sections in expandable containers
#                 section_config = [
#                     ("idea_evaluation", "💡 Startup Idea Evaluation"),
#                     ("uniqueness", "🔍 Uniqueness Check & Market Differentiation"),
#                     ("tech_stack", "⚙️ Technology Stack & Architecture"),
#                     ("similar_startups", "🏢 Similar Startups & Competition"),
#                     ("pitch", "💼 Pitch Generation - Professional Elevator Pitches"),
#                     ("improvements", "📈 Improvement Suggestions"),
#                     ("success_metrics", "💰 Success Probability & Investment Metrics"),
#                     ("additional", "📋 Additional Strategic Analysis")
#                 ]
                
#                 for section_key, title in section_config:
#                     content = evaluation_sections.get(section_key, "")
#                     if content.strip():
#                         with st.expander(title, expanded=False):
#                             st.markdown(content)
                
#                 # Metrics
#                 st.markdown("### 📊 Evaluation Metrics")
#                 evaluation_time = time.time() - start_time
                
#                 col1, col2, col3, col4 = st.columns(4)
#                 with col1:
#                     st.metric("🔥 API Calls", st.session_state.evaluator.api_call_count)
#                 with col2:
#                     st.metric("⏱️ Time", f"{evaluation_time:.1f}s")
#                 with col3:
#                     st.metric("📚 Sources", len(st.session_state.evaluator.conversation_memory[-1]['contexts']) if st.session_state.evaluator.conversation_memory else 0)
#                 with col4:
#                     st.metric("📁 Files", "Yes" if uploaded_content else "No")
                
#                 # Export
#                 st.markdown("### 💾 Export Report")
#                 col1, col2 = st.columns(2)
                
#                 with col1:
#                     full_report = f"""
# Founder's Fuel - STARTUP EVALUATION REPORT
# Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
# Startup: {startup_name or 'Not specified'}

# IDEA: {startup_idea}

# EVALUATION:
# {chr(10).join([f"{title}:{chr(10)}{content}{chr(10)}" for (_, title), content in zip(section_config, evaluation_sections.values()) if content.strip()])}

# METRICS:
# - API Calls: {st.session_state.evaluator.api_call_count}
# - Time: {evaluation_time:.1f}s
# - Session: {st.session_state.evaluator.user_session_id}
# """
                    
#                     st.download_button(
#                         "📄 Download Report",
#                         full_report,
#                         file_name=f"startwise_{startup_name or 'evaluation'}_{datetime.now().strftime('%Y%m%d_%H%M')}.txt",
#                         mime="text/plain"
#                     )
                
#                 with col2:
#                     json_report = {
#                         "startup_name": startup_name or "Not specified",
#                         "startup_idea": startup_idea,
#                         "evaluation_sections": evaluation_sections,
#                         "metadata": {
#                             "api_calls": st.session_state.evaluator.api_call_count,
#                             "time": evaluation_time,
#                             "timestamp": datetime.now().isoformat(),
#                             "uploaded_files": uploaded_files_info
#                         }
#                     }
                    
#                     st.download_button(
#                         "📊 Download JSON",
#                         json.dumps(json_report, indent=2),
#                         file_name=f"startwise_{startup_name or 'evaluation'}_{datetime.now().strftime('%Y%m%d_%H%M')}.json",
#                         mime="application/json"
#                     )
            
#             else:
#                 st.error("Evaluation failed. Please try again.")
        
#         elif submitted:
#             st.warning("Please enter your startup idea to begin evaluation.")

#     with tab2:
#         st.markdown("### 💬 Chat with Founder's Fuel")
        
#         # Initialize chat history
#         if 'chat_history' not in st.session_state:
#             st.session_state.chat_history = []
        
#         # Display chat
#         for msg in st.session_state.chat_history:
#             if msg['role'] == 'user':
#                 with st.chat_message("user"):
#                     st.write(msg['content'])
#             else:
#                 with st.chat_message("assistant"):
#                     st.write(msg['content'])
                    
#                     # Context info
#                     if msg.get('context_count', 0) > 0:
#                         st.info(f"📚 Used {msg['context_count']} knowledge sources")
        
#         # Chat input
#         user_input = st.chat_input("Ask about markets, competitors, tech stacks...")
        
#         if user_input:
#             st.session_state.chat_history.append({'role': 'user', 'content': user_input})
            
#             with st.spinner("Thinking..."):
#                 response, contexts = st.session_state.evaluator.enhanced_chat_with_web_context(user_input)
                
#                 st.session_state.chat_history.append({
#                     'role': 'assistant',
#                     'content': response,
#                     'context_count': len(contexts)
#                 })
                
#                 st.session_state.evaluator.save_conversation(user_input, response, len(contexts))
                
#                 # Fixed: Queue TTS instead of direct threading
#                 if st.session_state.get('tts_enabled', True):
#                     st.session_state.evaluator.text_to_speech(response)
            
#             st.rerun()

#     # Footer
#     st.markdown("---")
#     col1, col2, col3 = st.columns(3)
    
#     with col1:
#         st.markdown("**🔧 Tech Stack**")
#         st.caption("Streamlit + Gemini + Pinecone")
    
#     with col2:
#         st.markdown("**📊 Performance**") 
#         st.caption("Single API Call Optimization")
    
#     with col3:
#         st.markdown("**🌐 Data Sources**")
#         st.caption("RAG + Live Web Intelligence")

#     # Cleanup on exit
#     import atexit
#     def cleanup():
#         if 'evaluator' in st.session_state:
#             st.session_state.evaluator.cleanup_tts()
#     atexit.register(cleanup)


# if __name__ == "__main__":
#     main()



import streamlit as st
import os
from dotenv import load_dotenv
import io
import google.generativeai as genai
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
import PyPDF2
import docx
import pinecone
import numpy as np
from sentence_transformers import SentenceTransformer
import uuid
import time
from datetime import datetime
import json
import requests
from serpapi import GoogleSearch

import sqlite3
import re
import unicodedata
import hashlib
import functools
from typing import List, Dict, Any, Optional
import logging
from bs4 import BeautifulSoup
from urllib.parse import urlparse
import pyttsx3
import threading
import queue

# Configure logging with DEBUG level
logging.basicConfig(
    level=logging.DEBUG,  # Changed from WARNING to DEBUG
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(),  # Console output
        logging.FileHandler('foundersfuel_debug.log')  # File output
    ]
)
logger = logging.getLogger(__name__)



class OptimizedStartupRAGEvaluator:
    def __init__(self):
        self.embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
        self.embedding_dimension = 384
        self.gemini_model = None
        self.pinecone_index = None
        self.conversation_memory = []
        self.user_session_id = str(uuid.uuid4())
        self.api_call_count = 0
        self.cache = {}
        # Fixed TTS implementation
        self.tts_queue = queue.Queue()
        self.tts_thread_running = False
        self.tts_engine = None
        self.setup_models()
        self.setup_database()
        self.setup_tts()
        
    def setup_tts(self):
        """Initialize text-to-speech engine with thread safety"""
        try:
            # Initialize TTS engine once
            self.tts_engine = pyttsx3.init()
            # Set properties for better speech quality
            self.tts_engine.setProperty('rate', 180)  # Speed
            self.tts_engine.setProperty('volume', 0.8)  # Volume
            st.session_state.tts_enabled = True
            
            # Start TTS worker thread
            self.start_tts_worker()
            
        except Exception as e:
            logger.error(f"TTS setup error: {e}")
            st.session_state.tts_enabled = False
    
    def start_tts_worker(self):
        """Start the TTS worker thread"""
        if not self.tts_thread_running:
            self.tts_thread_running = True
            worker_thread = threading.Thread(target=self.tts_worker, daemon=True)
            worker_thread.start()
    
    def tts_worker(self):
        """TTS worker thread that processes speech requests"""
        while self.tts_thread_running:
            try:
                # Get text from queue with timeout
                text = self.tts_queue.get(timeout=1)
                if text is None:  # Shutdown signal
                    break
                
                # Clean text for TTS
                clean_text = re.sub(r'[#*\-•]', '', text)
                clean_text = re.sub(r'\s+', ' ', clean_text.strip())
                
                # Limit text length for TTS
                if len(clean_text) > 300:
                    clean_text = clean_text[:300] + "..."
                
                # Speak the text (this is thread-safe)
                if self.tts_engine and clean_text.strip():
                    self.tts_engine.say(clean_text)
                    self.tts_engine.runAndWait()
                
                self.tts_queue.task_done()
                
            except queue.Empty:
                continue
            except Exception as e:
                logger.error(f"TTS worker error: {e}")
                continue

    def setup_models(self):
        """Initialize Gemini and Pinecone"""
        try:
            # ---------------- GEMINI ----------------
            api_key = st.secrets.get("GEMINI_API_KEY")
            if not api_key:
                st.error("❌ GEMINI_API_KEY not found in Streamlit secrets")
                return
            genai.configure(api_key=api_key)
            self.gemini_model = genai.GenerativeModel('models/gemini-2.5-flash')

            # ---------------- PINECONE ----------------
            pinecone_api_key = st.secrets.get("PINECONE_API_KEY")
            pinecone_env = st.secrets.get("PINECONE_ENV", None)
            if not pinecone_api_key:
                st.error("❌ PINECONE_API_KEY not found in Streamlit secrets")
                return

            # Initialize Pinecone client
            if pinecone_env:
                pinecone.init(api_key=pinecone_api_key, environment=pinecone_env)
            else:
                pinecone.init(api_key=pinecone_api_key)

            index_name = "startwise-rag-knowledge"

            # Create index if it does not exist
            existing_indexes = pinecone.list_indexes()
            if index_name not in existing_indexes:
                st.write(f"Creating Pinecone index: **{index_name}**...")
                pinecone.create_index(
                    name=index_name,
                    dimension=self.embedding_dimension,
                    metric="cosine"
                )
                # small wait to let the service stabilize
                time.sleep(2)

            # Attach index client
            self.pinecone_index = pinecone.Index(index_name)
            logger.info(f"Pinecone index '{index_name}' is ready and attached.")

        except Exception as e:
            st.error(f"Setup error: {e}")

    def setup_database(self):
        """Enhanced database setup with pitch deck storage"""
        try:
            conn = sqlite3.connect('startup_evaluations.db', check_same_thread=False)
            cursor = conn.cursor()
            
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS evaluations (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    session_id TEXT,
                    startup_name TEXT,
                    startup_idea TEXT,
                    evaluation_results TEXT,
                    pitch_deck_content TEXT,
                    uploaded_files TEXT,
                    timestamp DATETIME DEFAULT CURRENT_TIMESTAMP
                )
            ''')
            
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS conversations (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    session_id TEXT,
                    user_message TEXT,
                    assistant_response TEXT,
                    context_used INTEGER,
                    timestamp DATETIME DEFAULT CURRENT_TIMESTAMP
                )
            ''')
            
            conn.commit()
            conn.close()
        except Exception as e:
            st.error(f"Database setup error: {e}")

    def text_to_speech(self, text: str):
        """Queue text for speech synthesis (thread-safe)"""
        try:
            if not st.session_state.get('tts_enabled', False) or not self.tts_engine:
                return
            
            # Add text to queue for processing
            self.tts_queue.put(text)
            
        except Exception as e:
            logger.error(f"TTS queue error: {e}")

    def auto_play_tts(self, text: str):
        """Safe auto-play TTS for chat responses"""
        self.text_to_speech(text)



    def check_knowledge_base_status(self) -> bool:
        """Check if the Pinecone index has vectors."""
        try:
            if not self.pinecone_index:
                st.warning("Pinecone index not initialized.")
                return False
            
            # describe_index_stats() is the key function to check the state
            stats = self.pinecone_index.describe_index_stats()
            vector_count = stats.get('total_vector_count', 0)
            
            logger.info(f"Pinecone index status: {vector_count} vectors found.")
            
            # If there are vectors, we assume the KB is loaded.
            return vector_count > 0
            
        except Exception as e:
            st.error(f"Failed to check knowledge base status: {e}")
            return False

    def clean_vector_id(self, vector_id: str) -> str:
        """Clean vector ID for ASCII compatibility"""
        cleaned = vector_id.replace('—', '-').replace('–', '-')
        cleaned = unicodedata.normalize('NFKD', cleaned)
        cleaned = cleaned.encode('ascii', 'ignore').decode('ascii')
        cleaned = re.sub(r'[^\w\-]', '_', cleaned)
        return cleaned.strip('_-')

    @functools.lru_cache(maxsize=200)
    def get_embedding(self, text: str):
        """Get embedding with caching"""
        return self.embedding_model.encode([text])[0]

    def extract_text(self, file_content: bytes, file_name: str, mime_type: str) -> str:
        """Enhanced text extraction with BeautifulSoup"""
        try:
            if mime_type == 'application/pdf':
                reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                text = "\n".join(page.extract_text() for page in reader.pages)
                
            elif 'wordprocessingml' in mime_type:
                doc = docx.Document(io.BytesIO(file_content))
                text = "\n".join(para.text for para in doc.paragraphs)
                
            elif mime_type == 'text/html':
                soup = BeautifulSoup(file_content, 'html.parser')
                for script in soup(["script", "style"]):
                    script.decompose()
                text = soup.get_text()
                
            else:
                text = file_content.decode('utf-8', errors='ignore')
            
            return re.sub(r'\s+', ' ', text.strip())
            
        except Exception as e:
            logger.error(f"Text extraction error: {e}")
            return ""

    def process_uploaded_files(self, uploaded_files) -> str:
        """Process uploaded pitch deck or startup documents"""
        combined_content = ""
        
        for uploaded_file in uploaded_files:
            try:
                file_content = uploaded_file.read()
                
                # Determine MIME type from file extension
                file_ext = uploaded_file.name.lower().split('.')[-1]
                mime_mapping = {
                    'pdf': 'application/pdf',
                    'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    'txt': 'text/plain',
                    'html': 'text/html'
                }
                
                mime_type = mime_mapping.get(file_ext, 'text/plain')
                
                # Extract text
                extracted_text = self.extract_text(file_content, uploaded_file.name, mime_type)
                
                if extracted_text:
                    combined_content += f"\n\n=== FROM {uploaded_file.name.upper()} ===\n"
                    combined_content += extracted_text[:2000]  # Limit per file
                    
            except Exception as e:
                st.warning(f"Could not process {uploaded_file.name}: {e}")
        
        return combined_content

    def web_search_with_scraping(self, query: str, max_results: int = 3) -> List[Dict]:
        """Enhanced web search with detailed logging"""
        try:
            logger.info(f"\n{'='*60}")
            logger.info(f"🌐 WEB SEARCH - Starting")
            logger.info(f"Query: {query}")
            logger.info(f"Max results: {max_results}")
            
            serpapi_key = st.secrets.get('SERPAPI_API_KEY')
            if not serpapi_key:
                logger.warning("⚠️ SERPAPI_API_KEY not found - skipping web search")
                return []

            search = GoogleSearch({
                "q": query,
                "api_key": serpapi_key,
                "engine": "google",
                "num": max_results
            })
            
            results = search.get_dict()
            organic_results = results.get("organic_results", [])
            
            logger.info(f"📊 Found {len(organic_results)} organic results")
            
            enhanced_results = []
            scraped_count = 0
            
            for idx, result in enumerate(organic_results):
                url = result.get('link', '')
                title = result.get('title', '')
                snippet = result.get('snippet', '')
                
                logger.debug(f"\n  Result #{idx+1}:")
                logger.debug(f"    - Title: {title}")
                logger.debug(f"    - URL: {url}")
                logger.debug(f"    - Snippet: {snippet[:100]}...")
                
                enhanced_result = {
                    'title': title,
                    'snippet': snippet,
                    'link': url,
                    'source': 'web_search',
                    'timestamp': datetime.now().isoformat()
                }
                
                # Check if URL should be scraped
                should_scrape = self.should_scrape_url(url)
                logger.debug(f"    - Should scrape: {should_scrape}")
                
                if should_scrape:
                    logger.info(f"  🔍 Attempting to scrape: {url}")
                    scraped_content = self.scrape_url_content(url)
                    if scraped_content:
                        enhanced_result['full_content'] = scraped_content
                        scraped_count += 1
                        logger.info(f"  ✅ Successfully scraped {len(scraped_content)} chars")
                    else:
                        logger.warning(f"  ⚠️ Scraping failed for {url}")
                
                enhanced_results.append(enhanced_result)
            
            logger.info(f"\n📈 WEB SEARCH SUMMARY:")
            logger.info(f"  - Total results: {len(enhanced_results)}")
            logger.info(f"  - Successfully scraped: {scraped_count}")
            logger.info(f"  - Results without scraping: {len(enhanced_results) - scraped_count}")
            logger.info(f"{'='*60}\n")
            
            return enhanced_results
            
        except Exception as e:
            logger.error(f"❌ Web search error: {e}", exc_info=True)
            return []

    def should_scrape_url(self, url: str) -> bool:
        """Determine if URL should be scraped"""
        if not url:
            return False
            
        trusted_domains = [
            'techcrunch.com', 'crunchbase.com', 'pitchbook.com', 
            'cbinsights.com', 'ycombinator.com', 'medium.com', 
            'forbes.com', 'entrepreneur.com', 'venturebeat.com',
            'bloomberg.com', 'reuters.com', 'wsj.com'
        ]
        
        parsed_url = urlparse(url)
        domain = parsed_url.netloc.lower()
        
        return any(trusted_domain in domain for trusted_domain in trusted_domains)

    def scrape_url_content(self, url: str) -> Optional[str]:
        """Scrape URL content using BeautifulSoup"""
        try:
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            }
            
            response = requests.get(url, headers=headers, timeout=10)
            response.raise_for_status()
            
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Remove unwanted elements
            for element in soup(['script', 'style', 'nav', 'footer', 'header']):
                element.decompose()
            
            # Extract main content
            main_content = soup.find('main') or soup.find('article') or soup.find('div', class_='content')
            
            if main_content:
                text = main_content.get_text(separator=' ', strip=True)
            else:
                text = soup.get_text(separator=' ', strip=True)
            
            # Clean the text
            text = re.sub(r'\s+', ' ', text)
            return text[:2000]  # Limit to prevent overload
            
        except Exception as e:
            logger.error(f"URL scraping error for {url}: {e}")
            return None

    def chunk_text(self, text: str, source_file: str, domain: str, chunk_size: int = 400) -> List[Dict]:
        """Simple text chunking"""
        words = text.split()
        chunks = []
        
        for i in range(0, len(words), chunk_size):
            chunk_words = words[i:i + chunk_size]
            chunk_text = ' '.join(chunk_words)
            
            if len(chunk_text.strip()) > 50:
                chunks.append({
                    'text': chunk_text,
                    'source_file': source_file,
                    'domain': domain,
                    'chunk_index': len(chunks),
                    'timestamp': datetime.now().isoformat()
                })
        
        return chunks

    def retrieve_context(self, query: str, top_k: int = 5) -> List[Dict]:
        """Enhanced context retrieval with detailed logging"""
        try:
            logger.info(f"\n{'='*60}")
            logger.info(f"🔍 RETRIEVE CONTEXT - Starting")
            logger.info(f"Query: {query[:100]}...")
            logger.info(f"Top K requested: {top_k}")
            
            query_embedding = self.get_embedding(query)
            logger.debug(f"Query embedding shape: {query_embedding.shape}")
            
            results = self.pinecone_index.query(
                vector=query_embedding.tolist(),
                top_k=top_k,
                include_metadata=True
            )
            
            total_matches = len(results.get('matches', []))
            logger.info(f"📊 Total matches found in Pinecone: {total_matches}")
            
            contexts = []
            filtered_out = 0
            
            for idx, match in enumerate(results.get('matches', [])):
                similarity = match['score']
                source = match['metadata'].get('source_file', 'Unknown')
                domain = match['metadata'].get('domain', 'Unknown')
                text_preview = match['metadata']['text'][:100]
                
                logger.debug(f"\n  Match #{idx+1}:")
                logger.debug(f"    - Similarity: {similarity:.4f}")
                logger.debug(f"    - Source: {source}")
                logger.debug(f"    - Domain: {domain}")
                logger.debug(f"    - Text preview: {text_preview}...")
                
                if match['score'] > 0.25:
                    contexts.append({
                        'text': match['metadata']['text'],
                        'source': source,
                        'domain': domain,
                        'similarity': similarity
                    })
                    logger.info(f"  ✅ Match #{idx+1} INCLUDED (similarity: {similarity:.4f})")
                else:
                    filtered_out += 1
                    logger.info(f"  ❌ Match #{idx+1} FILTERED OUT (similarity: {similarity:.4f} < 0.25)")
            
            logger.info(f"\n📈 RETRIEVAL SUMMARY:")
            logger.info(f"  - Total matches found: {total_matches}")
            logger.info(f"  - Matches included (>0.25): {len(contexts)}")
            logger.info(f"  - Matches filtered out: {filtered_out}")
            logger.info(f"{'='*60}\n")
            
            return contexts
            
        except Exception as e:
            logger.error(f"❌ Context retrieval error: {e}", exc_info=True)
            return []

    def comprehensive_startup_evaluation(self, startup_idea: str, startup_name: str = "", uploaded_content: str = "") -> Dict[str, str]:
        """Enhanced evaluation with comprehensive logging"""
        
        logger.info(f"\n{'#'*80}")
        logger.info(f"🚀 COMPREHENSIVE STARTUP EVALUATION - STARTING")
        logger.info(f"{'#'*80}")
        logger.info(f"Startup Name: {startup_name or 'Not specified'}")
        logger.info(f"Idea Length: {len(startup_idea)} chars")
        logger.info(f"Uploaded Content: {len(uploaded_content)} chars")
        
        st.write("🔍 Gathering knowledge base context...")
        logger.info("\n--- PHASE 1: RAG CONTEXT RETRIEVAL ---")
        rag_contexts = self.retrieve_context(startup_idea, top_k=8)
        logger.info(f"✅ Retrieved {len(rag_contexts)} RAG contexts")
        
        st.write("🌐 Searching for current market data...")
        logger.info("\n--- PHASE 2: WEB SEARCH & SCRAPING ---")
        
        web_searches = [
            f"{startup_idea} startup competitors market analysis",
            f"{startup_idea} technology stack architecture",
            f"{startup_idea} funding investment trends 2024"
        ]
        
        all_web_contexts = []
        for search_idx, search_query in enumerate(web_searches, 1):
            logger.info(f"\n🔎 Web Search #{search_idx}: {search_query}")
            web_results = self.web_search_with_scraping(search_query, max_results=2)
            all_web_contexts.extend(web_results)
            logger.info(f"  Retrieved {len(web_results)} results")
        
        logger.info(f"\n📊 Total web contexts collected: {len(all_web_contexts)}")
        
        # Build comprehensive context
        logger.info("\n--- PHASE 3: CONTEXT BUILDING ---")
        context_text = ""
        
        if uploaded_content.strip():
            logger.info(f"✅ Adding uploaded content: {len(uploaded_content)} chars")
            context_text += "\n=== UPLOADED STARTUP DOCUMENTS ===\n"
            context_text += uploaded_content[:2000] + "\n\n"
        
        if rag_contexts:
            logger.info(f"✅ Adding {len(rag_contexts[:5])} RAG contexts")
            context_text += "\n=== EXPERT KNOWLEDGE BASE ===\n"
            for ctx in rag_contexts[:5]:
                logger.debug(f"  - {ctx['domain']} | {ctx['source']} | {ctx['similarity']:.4f}")
                context_text += f"[{ctx['domain']} - {ctx['source']}] (Relevance: {ctx['similarity']:.2f})\n"
                context_text += f"{ctx['text']}\n\n"
        
        if all_web_contexts:
            logger.info(f"✅ Adding {len(all_web_contexts)} web contexts")
            context_text += "\n=== CURRENT MARKET INTELLIGENCE ===\n"
            for ctx in all_web_contexts:
                logger.debug(f"  - {ctx['title'][:50]}...")
                context_text += f"• {ctx['title']}\n{ctx['snippet']}\n"
                if ctx.get('full_content'):
                    context_text += f"Full Content: {ctx['full_content'][:400]}...\n\n"
        
        logger.info(f"\n📏 Total context length: {len(context_text)} chars")
        
        # Generate evaluation
        logger.info("\n--- PHASE 4: GEMINI API CALL ---")
        comprehensive_prompt = f"""You are FoundersFuelAI, an expert startup advisor. Analyze this startup covering ALL 7 critical evaluation domains.

STARTUP INFORMATION:
- Name: {startup_name or 'Not specified'}
- Idea: {startup_idea}

{context_text}

Provide a comprehensive evaluation covering these EXACT 7 domains:

## 1. STARTUP IDEA EVALUATION - Comprehensive Business Concept Analysis
- Problem-solution fit analysis
- Target market identification and size
- Business model viability assessment
- Value proposition strength
- Market timing and opportunity window
- Core concept validation score (1-10) with reasoning

## 2. UNIQUENESS CHECK - Market Differentiation & Competitive Analysis
- Analyze market differentiation potential
- Identify direct and indirect competitors
- Assess competitive advantages and moats
- Market positioning opportunities
- Uniqueness score (1-10) with reasoning

## 3. TECH STACK RECOMMENDATION - Technology Choices & Architecture
- Recommended technology stack (frontend, backend, database)
- Architecture patterns and scalability considerations
- Third-party integrations and APIs needed
- Development complexity assessment
- Technology risk factors

## 4. SIMILAR STARTUPS - Market Landscape & Competitors
- List 5-8 direct/indirect competitors with details
- Competitive landscape mapping
- Market share analysis
- Success/failure patterns in this space
- Competitive intelligence insights

## 5. PITCH GENERATION - Professional Elevator Pitches for Investors
### 60-Second Hooking Pitch:
[Create a compelling 60-second pitch designed to hook investor attention - focus on the problem, solution, and initial traction]

### In-Depth Investment Pitch:
[Comprehensive pitch with key points to elaborate:]
- Problem Statement & Market Pain
- Solution & Technology Advantage
- Market Size & Opportunity
- Business Model & Revenue Streams
- Competitive Landscape & Differentiation
- Traction & Metrics
- Team & Execution Capability
- Financial Projections & Unit Economics
- Funding Requirements & Use of Funds
- Exit Strategy & ROI Potential

### Key Messaging Framework:
- Value proposition statements
- Investment thesis points
- Risk mitigation talking points

## 6. IMPROVEMENT SUGGESTIONS - Enhancement Recommendations
- Product development priorities
- Market strategy optimizations
- Business model improvements
- Technical architecture enhancements
- Go-to-market refinements

## 7. SUCCESS PROBABILITY & INVESTMENT METRICS
- Overall success probability (%) with detailed reasoning
- Investment attractiveness rating (1-10)
- Market timing assessment
- Scalability potential
- Risk-reward analysis
- Funding readiness score

ADDITIONAL ANALYSIS:
- Revenue model recommendations
- Key performance indicators (KPIs)
- Milestone roadmap (6-month, 1-year, 3-year)
- Partnership opportunities
- Market entry strategy

Base your analysis on the provided expert knowledge, current market data, and uploaded documents. Be specific with numbers, examples, and actionable recommendations. Reference sources when making claims.
**IMPORTANT: Your entire response must be strictly under 3235 characters.**"""

        try:
            st.write("🧠 Generating comprehensive evaluation...")
            logger.info("📤 Sending request to Gemini API...")
            
            response = self.gemini_model.generate_content(comprehensive_prompt)
            self.api_call_count += 1
            
            logger.info(f"✅ Received response: {len(response.text)} chars")
            logger.info(f"📊 Total API calls so far: {self.api_call_count}")
            
            # Parse response
            logger.info("\n--- PHASE 5: RESPONSE PARSING ---")
            sections = self.parse_evaluation_response(response.text)
            
            for section_name, content in sections.items():
                if content.strip():
                    logger.debug(f"  ✅ Section '{section_name}': {len(content)} chars")
            
            logger.info(f"\n{'#'*80}")
            logger.info(f"✅ EVALUATION COMPLETE")
            logger.info(f"{'#'*80}\n")
            
            return sections
            
        except Exception as e:
            logger.error(f"❌ Error generating evaluation: {e}", exc_info=True)
            return {"error": f"Error generating evaluation: {e}"}

    def parse_evaluation_response(self, response_text: str) -> Dict[str, str]:
        """Parse the evaluation response into structured sections"""
        sections = {
            "idea_evaluation": "",
            "uniqueness": "",
            "tech_stack": "",
            "similar_startups": "",
            "pitch": "",
            "improvements": "",
            "success_metrics": "",
            "additional": ""
        }
        
        try:
            # Simple section parsing based on headers
            current_section = "general"
            lines = response_text.split('\n')
            
            for line in lines:
                lower_line = line.lower()
                
                if "startup idea evaluation" in lower_line or "business concept analysis" in lower_line:
                    current_section = "idea_evaluation"
                elif "uniqueness check" in lower_line or "market differentiation" in lower_line:
                    current_section = "uniqueness"
                elif "tech stack" in lower_line or "technology choices" in lower_line:
                    current_section = "tech_stack"
                elif "similar startups" in lower_line or "market landscape" in lower_line:
                    current_section = "similar_startups"
                elif "pitch generation" in lower_line or "elevator pitch" in lower_line:
                    current_section = "pitch"
                elif "improvement suggestions" in lower_line or "enhancement recommendations" in lower_line:
                    current_section = "improvements"
                elif "success probability" in lower_line or "investment metrics" in lower_line:
                    current_section = "success_metrics"
                elif "additional analysis" in lower_line:
                    current_section = "additional"
                else:
                    if current_section in sections:
                        sections[current_section] += line + "\n"
                    else:
                        sections["additional"] += line + "\n"
            
            # If parsing fails, put everything in additional
            if not any(sections.values()):
                sections["additional"] = response_text
                
        except Exception as e:
            logger.error(f"Parsing error: {e}")
            sections["additional"] = response_text
        
        return sections

    def enhanced_chat_with_web_context(self, user_message: str) -> tuple[str, List[Dict]]:
        """Enhanced chat with detailed logging"""
        
        logger.info(f"\n{'='*60}")
        logger.info(f"💬 CHAT REQUEST - Starting")
        logger.info(f"User message: {user_message[:100]}...")
        
        # Get RAG context
        logger.info("\n--- RAG Context Retrieval ---")
        rag_contexts = self.retrieve_context(user_message, top_k=5)
        logger.info(f"Retrieved {len(rag_contexts)} RAG contexts")
        
        # Check if web search needed
        web_contexts = []
        market_keywords = ['market', 'competitor', 'trend', 'industry', 'funding', 'investment']
        needs_web_search = any(keyword in user_message.lower() for keyword in market_keywords)
        
        logger.info(f"\n--- Web Search Check ---")
        logger.info(f"Needs web search: {needs_web_search}")
        
        if needs_web_search:
            logger.info("Performing web search...")
            web_contexts = self.web_search_with_scraping(user_message, max_results=3)
            logger.info(f"Retrieved {len(web_contexts)} web contexts")
        
        # Build context
        logger.info(f"\n--- Context Summary ---")
        logger.info(f"RAG contexts: {len(rag_contexts)}")
        logger.info(f"Web contexts: {len(web_contexts)}")
        logger.info(f"Conversation memory: {len(self.conversation_memory)} turns")
        
        context_text = ""
        
        if rag_contexts:
            context_text += "Knowledge Base:\n"
            for ctx in rag_contexts[:3]:
                logger.debug(f"  - {ctx['domain']}: {ctx['text'][:50]}...")
                context_text += f"[{ctx['domain']}] {ctx['text'][:300]}...\n\n"
        
        if web_contexts:
            context_text += "Current Market Data:\n"
            for ctx in web_contexts:
                logger.debug(f"  - {ctx['title'][:50]}...")
                context_text += f"• {ctx['title']}: {ctx['snippet']}\n"
                if ctx.get('full_content'):
                    context_text += f"Details: {ctx['full_content'][:200]}...\n\n"
        
        if self.conversation_memory:
            context_text += "\nRecent conversation:\n"
            for turn in self.conversation_memory[-2:]:
                context_text += f"Q: {turn['user']}\nA: {turn['assistant'][:200]}...\n\n"
        
        logger.info(f"Total context length: {len(context_text)} chars")
        
        # API call
        logger.info("\n--- Gemini API Call ---")
        prompt = f"""You are FoundersFuelAI, a startup advisor. Answer this question using the provided context.

{context_text}

User Question: {user_message}

Provide a helpful, specific answer based on the context above. Reference sources when relevant. **Your entire response must be under 3235 characters.**"""

        try:
            logger.info("Sending chat request to Gemini...")
            response = self.gemini_model.generate_content(prompt)
            self.api_call_count += 1
            
            logger.info(f"✅ Received response: {len(response.text)} chars")
            logger.info(f"Total API calls: {self.api_call_count}")
            
            # Save to memory
            self.conversation_memory.append({
                'user': user_message,
                'assistant': response.text,
                'contexts': rag_contexts + web_contexts
            })
            
            if len(self.conversation_memory) > 10:
                removed = len(self.conversation_memory) - 5
                self.conversation_memory = self.conversation_memory[-5:]
                logger.info(f"Trimmed {removed} old conversation turns")
            
            logger.info(f"{'='*60}\n")
            
            return response.text, rag_contexts + web_contexts
            
        except Exception as e:
            logger.error(f"❌ Chat error: {e}", exc_info=True)
            return f"Error: {e}", []

    def build_knowledge_base(self, service, parent_folder_id: str) -> bool:
        """Build knowledge base with progress tracking"""
        try:
            logger.info(f"\n{'#'*80}")
            logger.info(f"🔧 KNOWLEDGE BASE BUILD - Starting")
            logger.info(f"{'#'*80}")
            logger.info(f"Parent Folder ID: {parent_folder_id}")
            
            # Get domain folders
            query = f"'{parent_folder_id}' in parents and mimeType='application/vnd.google-apps.folder'"
            response = service.files().list(q=query, fields='files(id, name)').execute()
            domain_folders = response.get('files', [])
            
            if not domain_folders:
                logger.error("❌ No domain folders found")
                st.error("No domain folders found")
                return False
            
            logger.info(f"📁 Found {len(domain_folders)} domain folders")
            for folder in domain_folders:
                logger.info(f"  - {folder['name']}")
            
            progress_bar = st.progress(0)
            total_vectors = 0
            
            for folder_idx, folder in enumerate(domain_folders):
                domain_name = folder['name']
                folder_id = folder['id']
                namespace = f"domain_{self.clean_vector_id(domain_name)}"
                
                logger.info(f"\n--- Processing Domain: {domain_name} ---")
                st.write(f"Processing: {domain_name}")
                
                files = self.get_drive_files(service, folder_id)
                logger.info(f"📄 Found {len(files)} files in {domain_name}")
                
                all_chunks = []
                
                for file_idx, file in enumerate(files):
                    logger.debug(f"  Processing file {file_idx+1}/{len(files)}: {file['name']}")
                    file_content = self.download_file(service, file)
                    if file_content:
                        content = self.extract_text(file_content, file['name'], file['mimeType'])
                        if content.strip():
                            chunks = self.chunk_text(content, file['name'], domain_name)
                            all_chunks.extend(chunks)
                            logger.debug(f"    Created {len(chunks)} chunks")
                
                logger.info(f"✅ Total chunks for {domain_name}: {len(all_chunks)}")
                
                # Upload to Pinecone
                if all_chunks:
                    logger.info(f"📤 Uploading {len(all_chunks)} vectors to Pinecone...")
                    vectors_to_upsert = []
                    for chunk in all_chunks:
                        embedding = self.get_embedding(chunk['text'])
                        vector_id = f"{self.clean_vector_id(domain_name)}_{self.clean_vector_id(chunk['source_file'])}_{chunk['chunk_index']}"
                        
                        vectors_to_upsert.append({
                            'id': vector_id,
                            'values': embedding.tolist(),
                            'metadata': {
                                'text': chunk['text'],
                                'source_file': chunk['source_file'],
                                'domain': chunk['domain'],
                                'chunk_index': chunk['chunk_index'],
                                'timestamp': chunk['timestamp']
                            }
                        })
                    
                    # Upload in batches
                    batch_size = 100
                    for i in range(0, len(vectors_to_upsert), batch_size):
                        batch = vectors_to_upsert[i:i + batch_size]
                        self.pinecone_index.upsert(vectors=batch, namespace=namespace)
                        logger.debug(f"  Uploaded batch {i//batch_size + 1}/{(len(vectors_to_upsert)-1)//batch_size + 1}")
                    
                    total_vectors += len(vectors_to_upsert)
                    logger.info(f"✅ Uploaded {len(vectors_to_upsert)} vectors for {domain_name}")
                
                progress_bar.progress((folder_idx + 1) / len(domain_folders))
            
            logger.info(f"\n{'#'*80}")
            logger.info(f"✅ KNOWLEDGE BASE BUILD COMPLETE")
            logger.info(f"Total vectors uploaded: {total_vectors}")
            logger.info(f"{'#'*80}\n")
            
            st.success(f"Knowledge base built: {total_vectors} vectors")
            return True
            
        except Exception as e:
            logger.error(f"❌ Knowledge base build error: {e}", exc_info=True)
            st.error(f"Knowledge base build error: {e}")
            return False

    def get_drive_files(self, service, folder_id: str) -> List[Dict]:
        """Get files from Google Drive folder"""
        try:
            query = f"'{folder_id}' in parents"
            results = service.files().list(
                q=query,
                fields="files(id, name, mimeType)"
            ).execute()
            
            files = results.get('files', [])
            return [f for f in files if f['mimeType'] in [
                'application/pdf',
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                'application/vnd.google-apps.document',
                'text/plain',
                'text/html'
            ]]
            
        except Exception as e:
            logger.error(f"File retrieval error: {e}")
            return []

    def download_file(self, service, file: Dict) -> Optional[bytes]:
        """Download file from Google Drive"""
        try:
            if file['mimeType'] == 'application/vnd.google-apps.document':
                request = service.files().export_media(fileId=file['id'], mimeType='text/plain')
            else:
                request = service.files().get_media(fileId=file['id'])
            
            file_content = io.BytesIO()
            downloader = MediaIoBaseDownload(file_content, request)
            
            done = False
            while not done:
                status, done = downloader.next_chunk()
            
            file_content.seek(0)
            return file_content.read()
            
        except Exception as e:
            logger.error(f"Download error: {e}")
            return None

    def save_evaluation(self, startup_name: str, startup_idea: str, results: Dict[str, str], uploaded_files: str = ""):
        """Save evaluation to database with uploaded files info"""
        try:
            conn = sqlite3.connect('startup_evaluations.db', check_same_thread=False)
            cursor = conn.cursor()
            cursor.execute('''
                INSERT INTO evaluations (session_id, startup_name, startup_idea, evaluation_results, uploaded_files)
                VALUES (?, ?, ?, ?, ?)
            ''', (self.user_session_id, startup_name, startup_idea, json.dumps(results), uploaded_files))
            conn.commit()
            conn.close()
        except Exception as e:
            logger.error(f"Save error: {e}")

    def save_conversation(self, user_msg: str, assistant_response: str, context_count: int):
        """Save conversation to database"""
        try:
            conn = sqlite3.connect('startup_evaluations.db', check_same_thread=False)
            cursor = conn.cursor()
            cursor.execute('''
                INSERT INTO conversations (session_id, user_message, assistant_response, context_used)
                VALUES (?, ?, ?, ?)
            ''', (self.user_session_id, user_msg, assistant_response, context_count))
            conn.commit()
            conn.close()
        except Exception as e:
            logger.error(f"Conversation save error: {e}")

    def authenticate_google_drive(self):
        """Authenticate using Streamlit secrets JSON"""
        try:
            if "GDRIVE_SERVICE_ACCOUNT" not in st.secrets:
                st.error("❌ GDRIVE_SERVICE_ACCOUNT missing in Streamlit secrets")
                return None

            # NEW, FIXED CODE:
            service_json = st.secrets["GDRIVE_SERVICE_ACCOUNT"]

            # 1. Strip whitespace/newlines from the start/end.
            # 2. Replace any remaining literal newline characters (\n) with an empty string.
            #    This converts the multiline secret into a single, clean line of JSON.
            cleaned_json_string = service_json.strip().replace('\n', '')

            # Attempt to load the cleaned string
            service_info = json.loads(cleaned_json_string)

            credentials = service_account.Credentials.from_service_account_info(
                service_info,
                scopes=['https://www.googleapis.com/auth/drive.readonly']
            )

            return build('drive', 'v3', credentials=credentials)

        except Exception as e:
            st.error(f"Authentication failed: {e}")
            return None


    def load_evaluation_history(self) -> List[tuple]:
        """Load evaluation history"""
        try:
            conn = sqlite3.connect('startup_evaluations.db', check_same_thread=False)
            cursor = conn.cursor()
            cursor.execute('''
                SELECT startup_name, startup_idea, timestamp 
                FROM evaluations 
                ORDER BY timestamp DESC 
                LIMIT 10
            ''')
            results = cursor.fetchall()
            conn.close()
            return results
        except:
            return []

    def cleanup_tts(self):
        """Cleanup TTS resources on shutdown"""
        try:
            self.tts_thread_running = False
            if hasattr(self, 'tts_queue'):
                self.tts_queue.put(None)  # Shutdown signal
            if self.tts_engine:
                self.tts_engine.stop()
        except Exception as e:
            logger.error(f"TTS cleanup error: {e}")


def main():
    st.set_page_config(
        page_title="Founder's Fuel - Complete Startup Evaluator", 
        page_icon="🚀",
        layout="wide"
    )
        # --- Create Google Drive service-account.json from Streamlit secrets ---
  

    # Clean, simple styling
    st.markdown("""
    <style>
    .main-header {
        text-align: center;
        padding: 2rem;
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        border-radius: 15px;
        margin-bottom: 2rem;
        box-shadow: 0 8px 32px rgba(102, 126, 234, 0.3);
    }
    .status-success {
        background: #d4edda;
        color: #155724;
        border: 2px solid #28a745;
        border-radius: 8px;
        text-align: center;
        padding: 15px;
        font-weight: bold;
    }
    .status-warning {
        background: #fff3cd;
        color: #856404;
        border: 1px solid #ffeeba;
        border-radius: 8px;
        text-align: center;
        padding: 15px;
    }
    .metric-card {
        background: white;
        padding: 1rem;
        border-radius: 10px;
        border: 1px solid #e2e8f0;
        box-shadow: 0 2px 4px rgba(0, 0, 0, 0.1);
        text-align: center;
    }
        [data-testid="stExpander"] {
    border-radius: 10px;
    border: 1px solid #e2e8f0;
    transition: box-shadow 0.3s ease-in-out;
}

[data-testid="stExpander"]:hover {
    box-shadow: 0 4px 8px rgba(0, 0, 0, 0.1);
}
                        
[data-testid="stExpander"] summary p {
    font-size: 22px !important;
    font-weight: bold !important;
            color:#9EFCFF !important;
}
    .st-expander-content div {
        font-size: 19px !important;
    }
    </style>
    """, unsafe_allow_html=True)
    
    # Header
    st.markdown("""
    <div class="main-header">
        <h1>🚀 Founder's Fuel</h1>
        <h3>Complete Startup Evaluation Platform</h3>
        <p>7-Domain Analysis • Web-Enhanced RAG • Single API Call Optimization</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Initialize evaluator and check KB status on first run
    if 'evaluator' not in st.session_state:
        with st.spinner("Initializing evaluator and checking knowledge base..."):
            st.session_state.evaluator = OptimizedStartupRAGEvaluator()
            # Automatically check if the KB is already built in Pinecone
            st.session_state.kb_loaded = st.session_state.evaluator.check_knowledge_base_status()

    # Simple sidebar
    with st.sidebar:
        st.markdown("### 🛠️ Setup & Configuration")
        
        # TTS Toggle for chat only
        tts_enabled = st.toggle("🔊 Enable Chat TTS", value=st.session_state.get('tts_enabled', True))
        st.session_state.tts_enabled = tts_enabled
        
        if tts_enabled:
            st.success("🔊 Chat TTS Active")
        else:
            st.warning("🔇 Chat TTS Disabled")
        
        st.divider()
        
        parent_folder_id = st.text_input(
            "Google Drive Knowledge Folder ID", 
            value="1fZ2p0uvCKsGtqrofcVc-SeGxq0YCmUNT"
        )
        
        if st.button("🔧 Build / Rebuild Knowledge Base", type="primary"):
            service = st.session_state.evaluator.authenticate_google_drive()
            if service:
                with st.spinner("Building knowledge base..."):
                    if st.session_state.evaluator.build_knowledge_base(service, parent_folder_id):
                        st.session_state.kb_loaded = True
                        st.success("Knowledge base ready!")
                        st.rerun()
        
        # Status
        if st.session_state.kb_loaded:
            st.success("✅ Knowledge Base Ready")
        else:
            st.warning("⚠️ Knowledge Base Not Loaded")
        
        st.divider()
        
        # Metrics
        st.markdown("### 📈 Usage Metrics")
        col1, col2 = st.columns(2)
        with col1:
            st.metric("Gemini Calls", st.session_state.evaluator.api_call_count)
        with col2:
            st.metric("Session", st.session_state.evaluator.user_session_id[-6:])

    # Main content
    if not st.session_state.kb_loaded:
        st.info("Please build the knowledge base first using the sidebar to unlock all features.")
        return

    # Tabs - Reordered to put evaluation first
    tab1, tab2 = st.tabs(["📊 7-Domain Evaluation", "💬 AI Chat"])
    
    with tab1:
        st.markdown("### 📊 Complete 7-Domain Startup Evaluation")
        
        # Document Upload Section
        st.markdown("### 📁 Upload Supporting Documents")
        uploaded_files = st.file_uploader(
            "Upload pitch deck, business plan, or other startup documents",
            type=['pdf', 'docx', 'txt', 'html'],
            accept_multiple_files=True,
            help="Upload documents to provide better context for your evaluation"
        )
        
        uploaded_content = ""
        if uploaded_files:
            st.success(f"📄 {len(uploaded_files)} file(s) uploaded!")
            
            for file in uploaded_files:
                st.write(f"• **{file.name}** ({file.size:,} bytes)")
            
            with st.spinner("Processing files..."):
                uploaded_content = st.session_state.evaluator.process_uploaded_files(uploaded_files)
                if uploaded_content:
                    st.success("✅ Documents processed and ready for evaluation!")
        
        st.divider()
        
        with st.form("evaluation_form"):
            startup_name = st.text_input("🏢 Startup Name (optional)")
            startup_idea = st.text_area(
                "💡 Startup Idea & Description", 
                height=120,
                placeholder="Describe your startup concept, problem it solves, target market..."
            )
            
            submitted = st.form_submit_button("🚀 Generate Complete Evaluation", type="primary")
        
        if submitted and startup_idea.strip():
            start_time = time.time()
            
            with st.spinner("Generating comprehensive analysis..."):
                evaluation_sections = st.session_state.evaluator.comprehensive_startup_evaluation(
                    startup_idea, startup_name, uploaded_content
                )
            
            if "error" not in evaluation_sections:
                # Save results
                uploaded_files_info = ", ".join([f.name for f in uploaded_files]) if uploaded_files else ""
                st.session_state.evaluator.save_evaluation(
                    startup_name, startup_idea, evaluation_sections, uploaded_files_info
                )
                
                st.success("📋 Evaluation Complete!")
                
                # Display sections in expandable containers
                section_config = [
                    ("idea_evaluation", "💡 Startup Idea Evaluation"),
                    ("uniqueness", "🔍 Uniqueness Check & Market Differentiation"),
                    ("tech_stack", "⚙️ Technology Stack & Architecture"),
                    ("similar_startups", "🏢 Similar Startups & Competition"),
                    ("pitch", "💼 Pitch Generation - Professional Elevator Pitches"),
                    ("improvements", "📈 Improvement Suggestions"),
                    ("success_metrics", "💰 Success Probability & Investment Metrics"),
                    ("additional", "📋 Additional Strategic Analysis")
                ]
                
                for section_key, title in section_config:
                    content = evaluation_sections.get(section_key, "")
                    if content.strip():
                        with st.expander(title, expanded=False):
                            st.markdown(content)
                
                # Metrics
                st.markdown("### 📊 Evaluation Metrics")
                evaluation_time = time.time() - start_time
                
                col1, col2, col3, col4 = st.columns(4)
                with col1:
                    st.metric("🔥 API Calls", st.session_state.evaluator.api_call_count)
                with col2:
                    st.metric("⏱️ Time", f"{evaluation_time:.1f}s")
                with col3:
                    st.metric("📚 Sources", len(st.session_state.evaluator.conversation_memory[-1]['contexts']) if st.session_state.evaluator.conversation_memory else 0)
                with col4:
                    st.metric("📁 Files", "Yes" if uploaded_content else "No")
                
                # Export
                st.markdown("### 💾 Export Report")
                col1, col2 = st.columns(2)
                
                with col1:
                    full_report = f"""
Founder's Fuel - STARTUP EVALUATION REPORT
Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
Startup: {startup_name or 'Not specified'}

IDEA: {startup_idea}

EVALUATION:
{chr(10).join([f"{title}:{chr(10)}{content}{chr(10)}" for (_, title), content in zip(section_config, evaluation_sections.values()) if content.strip()])}

METRICS:
- API Calls: {st.session_state.evaluator.api_call_count}
- Time: {evaluation_time:.1f}s
- Session: {st.session_state.evaluator.user_session_id}
"""
                    
                    st.download_button(
                        "📄 Download Report",
                        full_report,
                        file_name=f"startwise_{startup_name or 'evaluation'}_{datetime.now().strftime('%Y%m%d_%H%M')}.txt",
                        mime="text/plain"
                    )
                
                with col2:
                    json_report = {
                        "startup_name": startup_name or "Not specified",
                        "startup_idea": startup_idea,
                        "evaluation_sections": evaluation_sections,
                        "metadata": {
                            "api_calls": st.session_state.evaluator.api_call_count,
                            "time": evaluation_time,
                            "timestamp": datetime.now().isoformat(),
                            "uploaded_files": uploaded_files_info
                        }
                    }
                    
                    st.download_button(
                        "📊 Download JSON",
                        json.dumps(json_report, indent=2),
                        file_name=f"startwise_{startup_name or 'evaluation'}_{datetime.now().strftime('%Y%m%d_%H%M')}.json",
                        mime="application/json"
                    )
            
            else:
                st.error("Evaluation failed. Please try again.")
        
        elif submitted:
            st.warning("Please enter your startup idea to begin evaluation.")

    with tab2:
        st.markdown("### 💬 Chat with Founder's Fuel")
        
        # Initialize chat history
        if 'chat_history' not in st.session_state:
            st.session_state.chat_history = []
        
        # Display chat
        for msg in st.session_state.chat_history:
            if msg['role'] == 'user':
                with st.chat_message("user"):
                    st.write(msg['content'])
            else:
                with st.chat_message("assistant"):
                    st.write(msg['content'])
                    
                    # Context info
                    if msg.get('context_count', 0) > 0:
                        st.info(f"📚 Used {msg['context_count']} knowledge sources")
        
        # Chat input
        user_input = st.chat_input("Ask about markets, competitors, tech stacks...")
        
        if user_input:
            st.session_state.chat_history.append({'role': 'user', 'content': user_input})
            
            with st.spinner("Thinking..."):
                response, contexts = st.session_state.evaluator.enhanced_chat_with_web_context(user_input)
                
                st.session_state.chat_history.append({
                    'role': 'assistant',
                    'content': response,
                    'context_count': len(contexts)
                })
                
                st.session_state.evaluator.save_conversation(user_input, response, len(contexts))
                
                # Fixed: Queue TTS instead of direct threading
                if st.session_state.get('tts_enabled', True):
                    st.session_state.evaluator.text_to_speech(response)
            
            st.rerun()

    # Footer
    st.markdown("---")
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.markdown("**🔧 Tech Stack**")
        st.caption("Streamlit + Gemini + Pinecone")
    
    with col2:
        st.markdown("**📊 Performance**") 
        st.caption("Single API Call Optimization")
    
    with col3:
        st.markdown("**🌐 Data Sources**")
        st.caption("RAG + Live Web Intelligence")

    # Cleanup on exit
    import atexit
    def cleanup():
        if 'evaluator' in st.session_state:
            st.session_state.evaluator.cleanup_tts()
    atexit.register(cleanup)


if __name__ == "__main__":
    main()